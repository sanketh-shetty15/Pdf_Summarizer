import os
import streamlit as st
from io import BytesIO
from docx import Document
import pdfkit

# For Q&A
import faiss
import numpy as np
from openai import OpenAI
from sentence_transformers import SentenceTransformer

# 👇 Absolute path to wkhtmltopdf.exe
path_wkhtmltopdf = r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"
config = pdfkit.configuration(wkhtmltopdf=path_wkhtmltopdf)

# Import your utility functions
from utils.pdf_utils import extract_text_from_pdf
from utils.summarizer_utils import summarize_text

# OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
embedder = SentenceTransformer("all-MiniLM-L6-v2")


# -------------------------------
# Helper: Chunking for Q&A
# -------------------------------
def chunk_text(text, max_chars=1000):
    """
    Split text into natural chunks (by paragraph) for better Q&A accuracy.
    """
    paragraphs = text.split("\n")
    chunks, current = [], ""

    for para in paragraphs:
        if len(current) + len(para) < max_chars:
            current += para + "\n"
        else:
            chunks.append(current.strip())
            current = para + "\n"

    if current.strip():
        chunks.append(current.strip())

    return chunks


# -------------------------------
# Streamlit UI
# -------------------------------
st.set_page_config(page_title="PDF Summariser + Q&A", page_icon="📄", layout="wide")
st.title("📄 AI PDF/Text Summariser + Q&A")

# Choose mode
mode = st.radio("Choose Mode:", ("Text", "PDF Summarizer", "PDF Q&A"))

# -------------------------------
# MODE 1: TEXT SUMMARIZER
# -------------------------------
if mode == "Text":
    raw_text = st.text_area("Paste your text here:", height=250)

    length_option = st.sidebar.selectbox("Summary Length", ["short", "medium", "long"])
    file_type = st.sidebar.selectbox("Download Format", ["pdf", "txt", "docx"], index=0)

    if st.button("Summarise"):
        if not raw_text.strip():
            st.warning("Please provide some input text.")
        else:
            with st.spinner("Summarising..."):
                summary = summarize_text(raw_text, length=length_option)

            st.subheader("📝 Summary")
            st.write(summary)

            # Prepare file for download
            file_data, file_name, mime_type = None, f"summary.{file_type}", "text/plain"

            if file_type == "txt":
                clean_summary = summary.replace("*", "").replace("#", "")
                file_data = clean_summary.encode("utf-8")
                mime_type = "text/plain"

            elif file_type == "docx":
                buffer = BytesIO()
                doc = Document()
                for line in summary.split("\n"):
                    clean_line = line.replace("*", "").replace("#", "").strip()
                    if not clean_line:
                        doc.add_paragraph("")
                    elif line.strip().startswith("-"):
                        doc.add_paragraph(clean_line.lstrip("-").strip(), style="List Bullet")
                    else:
                        if (
                            clean_line.isupper()
                            or clean_line[0].isdigit()
                            or (len(clean_line) > 1 and clean_line[0] in "IVXLCDM" and clean_line[1] == ".")
                        ):
                            doc.add_paragraph(clean_line, style="Heading2")
                        else:
                            doc.add_paragraph(clean_line)
                doc.save(buffer)
                buffer.seek(0)
                file_data = buffer.getvalue()
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

            elif file_type == "pdf":
                html_content = """
                <html>
                <head>
                    <style>
                        body { font-family: 'Segoe UI', Arial, sans-serif; line-height: 1.6; margin: 40px; color: #2C3E50; }
                        h1 { text-align: center; color: #1A5276; font-size: 28px; margin-bottom: 30px; font-weight: bold; }
                        h2 { color: #2E86C1; border-bottom: 2px solid #D6EAF8; padding-bottom: 4px; margin-top: 24px; font-size: 20px; font-weight: bold; }
                        p { margin: 6px 0; font-size: 13pt; }
                        ul { margin: 8px 0 8px 20px; font-size: 13pt; }
                    </style>
                </head>
                <body>
                    <h1>Summary</h1>
                """
                for line in summary.split("\n"):
                    stripped = line.replace("*", "").replace("#", "").strip()
                    if not stripped:
                        html_content += "<br/>"
                    elif line.strip().startswith("-"):
                        clean_bullet = stripped.lstrip("-").strip()
                        html_content += f"<ul><li>{clean_bullet}</li></ul>"
                    elif (
                        stripped.isupper()
                        or stripped[0].isdigit()
                        or (len(stripped) > 1 and stripped[0] in "IVXLCDM" and stripped[1] == ".")
                    ):
                        html_content += f"<h2>{stripped}</h2>"
                    else:
                        html_content += f"<p>{stripped}</p>"

                html_content += "</body></html>"

                pdf_bytes = pdfkit.from_string(html_content, False, configuration=config)
                file_data = pdf_bytes
                mime_type = "application/pdf"

            st.download_button(
                "⬇️ Download Summary",
                data=file_data,
                file_name=file_name,
                mime=mime_type
            )

# -------------------------------
# MODE 2: PDF SUMMARIZER
# -------------------------------
elif mode == "PDF Summarizer":
    uploaded_file = st.file_uploader("Upload a PDF", type="pdf")
    if uploaded_file:
        raw_text = extract_text_from_pdf(uploaded_file)

        length_option = st.sidebar.selectbox("Summary Length", ["short", "medium", "long"])
        file_type = st.sidebar.selectbox("Download Format", ["pdf", "txt", "docx"], index=0)

        if st.button("Summarise"):
            if not raw_text.strip():
                st.warning("Please upload a valid PDF.")
            else:
                with st.spinner("Summarising..."):
                    summary = summarize_text(raw_text, length=length_option)

                st.subheader("📝 Summary")
                st.write(summary)

# -------------------------------
# MODE 3: PDF Q&A
# -------------------------------
elif mode == "PDF Q&A":
    uploaded_file = st.file_uploader("Upload a PDF", type="pdf")
    if uploaded_file:
        raw_text = extract_text_from_pdf(uploaded_file)

        # Use paragraph-based chunking
        chunks = chunk_text(raw_text, max_chars=1000)
        embeddings = embedder.encode(chunks, convert_to_numpy=True)
        dim = embeddings.shape[1]
        index = faiss.IndexFlatL2(dim)
        index.add(embeddings)

        st.success("✅ PDF processed! Now ask your questions below 👇")

        user_q = st.text_input("Ask a question about the PDF:")
        if user_q:
            q_embedding = embedder.encode([user_q], convert_to_numpy=True)
            D, I = index.search(q_embedding, 5)  # top-5 chunks
            retrieved = [chunks[i] for i in I[0]]

            context = "\n\n".join(retrieved)

            strict_prompt = f"""
You are an assistant that ONLY answers using the provided context.
If the answer is not in the context, reply with "I cannot find the answer in the document."

Context:
{context}

Question: {user_q}
Answer:
"""

            with st.spinner("Finding answer..."):
                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "You are a strict PDF assistant."},
                        {"role": "user", "content": strict_prompt}
                    ]
                )

            st.subheader("📌 Answer")
            st.write(response.choices[0].message.content.strip())

            with st.expander("🔎 Context used for this answer"):
                st.write(context)
